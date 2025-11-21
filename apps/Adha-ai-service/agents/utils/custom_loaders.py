"""
Custom document loaders to replace LangChain loaders.
Lightweight alternatives using PyMuPDF (already installed) and python-docx.
"""
import os
from typing import List, Dict
import fitz  # PyMuPDF


class Document:
    """Simple document container compatible with our RAG system."""
    def __init__(self, page_content: str, metadata: Dict = None):
        self.page_content = page_content
        self.metadata = metadata or {}


class PyPDFLoader:
    """PDF loader using PyMuPDF (fitz) - already installed."""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """Load PDF and return list of Document objects (one per page)."""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"PDF file not found: {self.file_path}")
        
        documents = []
        try:
            pdf_document = fitz.open(self.file_path)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text = page.get_text()
                
                doc = Document(
                    page_content=text.strip(),
                    metadata={
                        "source": self.file_path,
                        "page": page_num + 1,
                        "total_pages": len(pdf_document)
                    }
                )
                documents.append(doc)
            
            pdf_document.close()
            return documents
            
        except Exception as e:
            raise Exception(f"Error loading PDF {self.file_path}: {str(e)}")


class TextLoader:
    """Simple text file loader."""
    
    def __init__(self, file_path: str, encoding: str = 'utf-8'):
        self.file_path = file_path
        self.encoding = encoding
    
    def load(self) -> List[Document]:
        """Load text file and return as single Document."""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"Text file not found: {self.file_path}")
        
        try:
            with open(self.file_path, 'r', encoding=self.encoding) as f:
                text = f.read()
            
            doc = Document(
                page_content=text.strip(),
                metadata={
                    "source": self.file_path
                }
            )
            return [doc]
            
        except Exception as e:
            raise Exception(f"Error loading text file {self.file_path}: {str(e)}")


class Docx2txtLoader:
    """DOCX loader using python-docx library."""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """Load DOCX file and return as single Document."""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"DOCX file not found: {self.file_path}")
        
        try:
            # Try to import python-docx
            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise ImportError(
                    "python-docx is required for DOCX loading. "
                    "Install it with: pip install python-docx"
                )
            
            docx = DocxDocument(self.file_path)
            
            # Extract text from all paragraphs
            text = "\n".join([paragraph.text for paragraph in docx.paragraphs])
            
            doc = Document(
                page_content=text.strip(),
                metadata={
                    "source": self.file_path,
                    "paragraphs": len(docx.paragraphs)
                }
            )
            return [doc]
            
        except Exception as e:
            raise Exception(f"Error loading DOCX {self.file_path}: {str(e)}")
